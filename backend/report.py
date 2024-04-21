from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from backend.system import secure_filename


def generate_report_text(filename: str, emotion: str) -> str:
    """
    Генерирует динамический текст отчета

    :param filename: Имя обработанного файла
    :param emotion: Эмоция, найденная в файле
    :return: Текст отчета
    """
    return f"В резуальтате обработке файла {filename}, был получен результат средней эмоции: {emotion}."


def create_report(text: str, filename: str, image_path: str | None) -> str:
    """
    Создает отчет в формате docx
    :param text: текст отчета
    :param filename: имя файла
    :return: Путь до отчета
    """
    doc = Document()

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(filename)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"

    # Текст
    text = doc.add_paragraph()
    text.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text_run = text.add_run(text)
    text_run.font.size = Pt(12)
    text_run.font.name = "Times New Roman"

    if image_path is not None:
        # Изображение
        image = doc.add_paragraph()
        image_run = image.add_run(image_path, width=Inches(3))
        image_run.width = Pt(300)

    # Сохранить
    filename = f"{filename[:filename.find('.')]}_report.docx"
    report_path = f"report/{filename}"
    doc.save(report_path)

    return filename
